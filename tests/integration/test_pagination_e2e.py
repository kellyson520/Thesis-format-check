import io
import os
import sys
import pytest
import docx

# 配置项目路径以便导入 src 模块
sys.path.append(os.path.join(os.getcwd(), "src"))

from infrastructure.parsers.docx_parser import DocxParser
from use_cases.rule_config import RuleLoader, RuleConfig
from use_cases.validator_pipeline import ValidatorPipeline

def create_p2_test_docx():
    doc = docx.Document()
    
    # 1. 标题 1：显式关闭 keep_with_next 和 widow_control
    p1 = doc.add_paragraph("第一章 绪论", style='Heading 1')
    p1.paragraph_format.keep_with_next = False
    p1.paragraph_format.widow_control = False
    
    # 2. 正文：显式关闭 widow_control
    p2 = doc.add_paragraph("这是一段正文。它可能会产生孤行，因为我们没有开启孤行控制。")
    p2.paragraph_format.widow_control = False
    
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream

def test_pagination_validation():
    # 1. 准备强类型规则
    rules_data = {
        "document": {"default_font_size": 12},
        "pagination": {
            "widow_control": True,
            "keep_with_next_styles": ["Heading 1", "标题 1"]
        }
    }
    config = RuleConfig.model_validate(rules_data)
    
    # 2. 准备文档并解析（Parser 内部会自动触发 RevisionCleaner）
    stream = create_p2_test_docx()
    parser = DocxParser(stream)
    doc_model = parser.parse()
    
    # 验证 Parser 属性提取
    assert doc_model.paragraphs[0].keep_with_next is False
    assert doc_model.paragraphs[0].widow_control is False
    assert doc_model.paragraphs[1].widow_control is False
    
    # 3. 运行校验流水线
    pipeline = ValidatorPipeline(config)
    all_issues = []
    for event in pipeline.run(doc_model):
        all_issues.extend(event.issues)
    
    # 4. 断言
    codes = [i.issue_code.value for i in all_issues]
    print(f"检测到的问题代码: {codes}")
    
    # 应该发现 E009 (孤行) 和 E010 (与下段同页)
    assert "E009" in codes # 正文段落缺失孤行控制
    assert "E010" in codes # 标题段落缺失与下段同页
    
    print("Pagination validation test passed!")

if __name__ == "__main__":
    test_pagination_validation()
    print("SUCCESS")