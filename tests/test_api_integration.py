import io
import docx
from fastapi.testclient import TestClient
import sys
import os

# Add src to Python Path
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '../src')))

from main import app, _API_TOKEN
import json

def test_full_api_loop():
    client = TestClient(app)
    headers = {"X-Token": _API_TOKEN}
    
    # 1. Create a dummy Docx
    doc = docx.Document()
    p = doc.add_paragraph("This is a test paragraph with wrong size.")
    run = p.add_run("Bad Size Run")
    run.font.size = docx.shared.Pt(8)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    content = buffer.getvalue()
    
    # 2. Test /api/check (Zero-Disk IO)
    response = client.post(
        "/api/check",
        headers=headers,
        files={"file": ("test.docx", content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    )
    assert response.status_code == 200
    data = response.json()
    assert "issues" in data
    print(f"Check success: {data['total']} issues found.")
    
    # 3. Test /api/check/stream (SSE)
    response_stream = client.post(
        "/api/check/stream",
        headers=headers,
        files={"file": ("test.docx", content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    )
    assert response_stream.status_code == 200
    # Collect events
    events = []
    for line in response_stream.iter_lines():
        if line.startswith("data: "):
            events.append(json.loads(line[6:]))
    
    # Verify events
    assert any(e.get("event_type") == "progress" for e in events)
    assert any(e.get("event_type") == "done" for e in events)
    # Check issue structure in stream
    progress_event = next(e for e in events if e.get("event_type") == "progress")
    if progress_event.get("issues"):
        assert "type" in progress_event["issues"][0]
    
    print("Stream success: SSE events validated.")

    # 4. Test /api/fix (Zero-Disk IO + FixerPipeline)
    response_fix = client.post(
        "/api/fix",
        headers=headers,
        files={"file": ("test.docx", content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    )
    assert response_fix.status_code == 200
    assert response_fix.headers["Content-Type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    
    print("Fix success: Received fixed docx.")

if __name__ == "__main__":
    try:
        test_full_api_loop()
        print("Integration Test Passed!")
    except Exception as e:
        print(f"Integration Test Failed: {e}")
        import traceback
        traceback.print_exc()
