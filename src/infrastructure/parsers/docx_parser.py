"""
Infrastructure Layer — DocxParser（纯翻译官）
职责：将 python-docx 的 XML DOM 翻译成 Domain Layer 的抽象文档模型（ADM）。
规则：
  - 不含任何业务逻辑判断（"正确"的概念属于 Plugin 层）。
  - 不调用 FastAPI / WebSocket。
  - 所有单位换算通过 UnitConverter 完成，禁止硬编码。
  - 实现 IParser 接口，use_cases 层对此文件一无所知。

P0 重构要点（2026-04-03）：
  - 接入 StyleResolver：当 run.bold / font / size 为 None 时，
    通过样式继承链获取真实计算值，消除误报。
  - 移除 _parse_runs() 中粗暴的 `if not run.text.strip(): continue`，
    保留纯空白 Run 的格式信息（首行缩进检测所需）。
  - 扩展解析：_parse_tables() 将表格单元格段落纳入 ADM。
"""
from __future__ import annotations
import docx
import io
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from domain.models import ParagraphNode, TextRunNode, SectionNode, DocumentModel
from domain.utils.unit_converter import UnitConverter
from domain.interfaces import IParser
from infrastructure.parsers.style_resolver import StyleResolver
from infrastructure.parsers.revision_cleaner import RevisionCleaner

# Word XML 命名空间
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_QNAME_RFONTS = f"{{{_W_NS}}}rFonts"
_ATTR_EAST_ASIA = f"{{{_W_NS}}}eastAsia"
_ATTR_ASCII     = f"{{{_W_NS}}}ascii"

# 对齐映射（从 WD_PARAGRAPH_ALIGNMENT 到统一字符串）
_ALIGN_MAP = {
    WD_PARAGRAPH_ALIGNMENT.CENTER:  "center",
    WD_PARAGRAPH_ALIGNMENT.LEFT:    "left",
    WD_PARAGRAPH_ALIGNMENT.RIGHT:   "right",
    WD_PARAGRAPH_ALIGNMENT.JUSTIFY: "justify",
}


def _get_rfonts_attr(run, attr: str):
    """从 Run XML 中提取 rFonts 属性（如 eastAsia / ascii）。"""
    try:
        rPr = run._element.rPr
        if rPr is not None:
            rFonts = rPr.find(_QNAME_RFONTS)
            if rFonts is not None:
                return rFonts.get(attr)
    except AttributeError:
        pass
    return None


class DocxParser(IParser):
    """
    将 .docx 文件翻译为 DocumentModel（ADM）。
    实现 IParser 接口，use_cases 层通过接口调用，不直接依赖此类。
    """

    def __init__(self, doc_stream: io.BytesIO, clean_revisions: bool = True):
        # python-docx 对象只在此类内部存在，不泄露到其他层
        self._doc = docx.Document(doc_stream)
        
        # ── P2: 开启修订模式屏蔽与垃圾清理 ──────────────────────────────────
        if clean_revisions:
            RevisionCleaner.accept_all_revisions(self._doc)
            RevisionCleaner.strip_hidden_text(self._doc)
            
        # P0 新增：预构建 StyleResolver，供 _parse_runs 回退查询
        self._resolver = StyleResolver(self._doc)

    def parse(self) -> DocumentModel:
        """执行解析，返回纯 Python 数据模型（无 docx 对象）。"""
        paragraphs = self._parse_paragraphs()
        table_paras = self._parse_tables()
        sections   = self._parse_sections()
        # 将表格段落追加至后端，携带 source 标记（暂存为全局段落索引延续）
        return DocumentModel(
            paragraphs=paragraphs + table_paras,
            sections=sections,
        )

    # ── 段落解析 ────────────────────────────────────────────────────────────

    def _parse_paragraphs(self) -> list[ParagraphNode]:
        result = []
        for i, para in enumerate(self._doc.paragraphs):
            node = self._para_to_node(para, global_index=i + 1)
            if node is not None:
                result.append(node)
        return result

    def _para_to_node(self, para, global_index: int) -> ParagraphNode | None:
        """
        将单个 python-docx Paragraph 转换为 ParagraphNode。
        空段落（无文本也无格式意义）返回 None 跳过。
        """
        # 保留含有格式但纯空白的段落（首行缩进检测需要）
        # 但对于完全无内容且无样式信息的段落，静默跳过
        text = para.text
        stripped = text.strip()
        if not stripped and not para.runs:
            return None

        style_name = para.style.name if para.style else "Normal"
        pf = para.paragraph_format

        # 获取段落继承字号，用于精确换算“字符宽度”
        para_font_size = self._resolver.get_computed_size_pt(style_name) or 12.0

        # 单位换算统一交给 UnitConverter
        indent_chars = UnitConverter.pt_to_chars(
            pf.first_line_indent.pt if (pf and pf.first_line_indent) else 0.0,
            para_font_size
        )
        line_spacing = UnitConverter.normalize_line_spacing(pf.line_spacing)
        space_before = pf.space_before.pt if pf and pf.space_before else 0.0
        space_after  = pf.space_after.pt  if pf and pf.space_after  else 0.0
        alignment    = _ALIGN_MAP.get(para.alignment)

        runs = self._parse_runs(para, style_name)

        # ── 分页属性解析 ────────────────────────────────────────────────────────
        widow_control  = pf.widow_control  if (pf and pf.widow_control  is not None) else self._resolver.get_computed_widow_control(style_name)
        keep_with_next = pf.keep_with_next if (pf and pf.keep_with_next is not None) else self._resolver.get_computed_keep_with_next(style_name)
        keep_together  = pf.keep_together  if (pf and pf.keep_together  is not None) else self._resolver.get_computed_keep_together(style_name)

        # 若段落没有有效文字内容（仅空白 run），且不是标题/题注样式则跳过
        if not stripped:
            is_structural = (
                style_name.startswith("Heading") or
                style_name.startswith("标题") or
                "Caption" in style_name or
                "题注" in style_name
            )
            if not is_structural:
                return None

        return ParagraphNode(
            index=global_index,
            text=stripped or text,
            style_name=style_name,
            alignment=alignment,
            line_spacing=line_spacing,
            first_line_indent_chars=indent_chars,
            space_before_pt=space_before,
            space_after_pt=space_after,
            widow_control=widow_control,
            keep_with_next=keep_with_next,
            keep_together=keep_together,
            runs=runs,
        )

    def _parse_runs(self, para, style_name: str) -> list[TextRunNode]:
        """
        解析段落内所有 Run。
        P1 变更：保留所有 Run（含空 Run），确保 run_index 能够与物理文档对齐。
        """
        result = []
        for run in para.runs:
            # ── 字体解析（优先 Run 直接格式，回退顺序：Run → 样式继承链）──
            ascii_font = (
                run.font.name
                or _get_rfonts_attr(run, _ATTR_ASCII)
                or self._resolver.get_computed_ascii_font(style_name)
            )
            ea_font = (
                _get_rfonts_attr(run, _ATTR_EAST_ASIA)
                or self._resolver.get_computed_east_asia_font(style_name)
            )
            if ascii_font is None and para.style and para.style.font:
                ascii_font = para.style.font.name

            # ── 字号解析（Run → 段落样式 → 继承链）──────────────────────
            size_pt: float | None = None
            if run.font.size:
                size_pt = run.font.size.pt
            elif para.style and para.style.font and para.style.font.size:
                size_pt = para.style.font.size.pt
            else:
                # P0 新增：回退到 StyleResolver 的样式继承值
                size_pt = self._resolver.get_computed_size_pt(style_name)

            # ── 字重解析（核心修复点）────────────────────────────────────
            # run.bold 有三种状态：True（加粗）/ False（明确不加粗）/ None（继承）
            # 旧代码：`run.bold if run.bold is not None else False`（错误：把继承误判为 False）
            # 新代码：None 时向 StyleResolver 查询继承链的真实渲染值
            if run.bold is not None:
                bold = run.bold
            else:
                inherited = self._resolver.get_computed_bold(style_name)
                bold = inherited if inherited is not None else False

            result.append(TextRunNode(
                text=run.text,
                ascii_font=ascii_font,
                east_asia_font=ea_font,
                size_pt=size_pt,
                bold=bold,
            ))
        return result

    # ── 表格解析（P0 新增）──────────────────────────────────────────────────

    def _parse_tables(self) -> list[ParagraphNode]:
        """
        遍历文档中所有表格的单元格段落，纳入 ADM。
        索引从 len(doc.paragraphs) + 1 开始，避免与正文段落索引冲突。
        """
        result = []
        base_index = len(self._doc.paragraphs) + 1

        for table in self._doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for j, para in enumerate(cell.paragraphs):
                        node = self._para_to_node(para, global_index=base_index)
                        if node is not None:
                            result.append(node)
                        base_index += 1

        return result

    # ── 节解析（页边距）────────────────────────────────────────────────────

    def _parse_sections(self) -> list[SectionNode]:
        result = []
        for s in self._doc.sections:
            result.append(SectionNode(
                top_margin_cm    = s.top_margin.cm    if s.top_margin    else None,
                bottom_margin_cm = s.bottom_margin.cm if s.bottom_margin else None,
                left_margin_cm   = s.left_margin.cm   if s.left_margin   else None,
                right_margin_cm  = s.right_margin.cm  if s.right_margin  else None,
            ))
        return result
