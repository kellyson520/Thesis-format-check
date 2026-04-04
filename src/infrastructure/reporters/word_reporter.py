"""
Infrastructure Layer — Word 报告生成器（批注注入 + 格式修复）
职责：
  1. AnnotationReporter — 将 Issue 列表以红色批注形式注入 Word 副本
  2. DocumentFixer      — 一键自动修复格式（覆盖段落样式）
规则：
  - 使用新 RuleConfig（Pydantic 强类型），不再使用裸字典
  - 接受 List[Issue] (domain.models.Issue)，不依赖 FastAPI
"""
from __future__ import annotations
import io
from typing import List

import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

from domain.models import Issue
from use_cases.rule_config import RuleConfig


# ─── 对齐映射（字符串 → WD_PARAGRAPH_ALIGNMENT） ──────────────────────────────
_ALIGN_REVERSE = {
    "center":  WD_PARAGRAPH_ALIGNMENT.CENTER,
    "left":    WD_PARAGRAPH_ALIGNMENT.LEFT,
    "right":   WD_PARAGRAPH_ALIGNMENT.RIGHT,
    "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
}


# ─── 1. 批注报告生成器 ────────────────────────────────────────────────────────

class AnnotationReporter:
    """
    将校验 Issue 以红色内联批注形式注入 Word 文档副本，返回字节流。
    接受 domain.models.Issue 列表（非裸字典）。
    """

    def __init__(self, doc_source: io.BytesIO | str):
        self._doc = docx.Document(doc_source)

    def generate(self, issues: List[Issue]) -> bytes:
        """
        生成带批注的 Word 文件。
        :param issues: ValidatorPipeline 产出的 Issue 列表
        :return: 修改后的 .docx 字节流
        """
        # 使用已加载的 doc
        doc = self._doc

        # 按段落索引聚合 issues
        issue_map: dict[int, list[Issue]] = {}
        for issue in issues:
            issue_map.setdefault(issue.para_index, []).append(issue)

        for para_idx, para in enumerate(doc.paragraphs, start=1):
            if para_idx not in issue_map:
                continue
            msgs = " | ".join(
                f"[{i.severity.value}] {i.message}"
                for i in issue_map[para_idx]
            )
            run = para.add_run(f"  ⚠ {msgs}")
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            run.font.size = Pt(9)
            run.italic = True

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()


