"""
Infrastructure Layer — DocxFixer
职责：将 domain.models.Patch 应用到物理的 python-docx 文档并输出二进制流。
规则：
  - 此类是 IFixer 接口的具体实现。
  - 使用 Zero-Disk IO（io.BytesIO）。
  - 处理物理 Word 细节（如 XML 命名空间设置中文字体）。
"""
from __future__ import annotations
import io
from typing import List, Optional
import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from domain.models import Patch
from domain.interfaces import IFixer

# 命名空间定义
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_QNAME_RFONTS = f"{{{_W_NS}}}rFonts"
_ATTR_EAST_ASIA = f"{{{_W_NS}}}eastAsia"


class DocxFixer(IFixer):
    """
    负责 .docx 文件的自动修复逻辑。
    """

    def __init__(self, doc_source: str | io.BytesIO):
        """
        :param doc_source: 文件路径或 BytesIO 缓存。
        """
        self._doc = docx.Document(doc_source)
        # 获取扁平化的文档段落引用，以对齐 DocxParser 的索引逻辑
        self._all_paras = self._get_flat_paragraphs()

    def _get_flat_paragraphs(self) -> List[docx.text.paragraph.Paragraph]:
        """
        获取文档中所有段落的引用（正文+表格内容）。
        顺序必须与 DocxParser._parse_paragraphs + _parse_tables 严格一致。
        """
        paras = list(self._doc.paragraphs)
        for table in self._doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paras.extend(cell.paragraphs)
        return paras

    def fix(self, patches: List[Patch]) -> bytes:
        """
        执行修复并产出 bytes 流。
        """
        for patch in patches:
            try:
                self._apply_patch(patch)
            except Exception as e:
                # 记录失败但继续处理其他补丁（TODO: 接入 AppLogger）
                print(f"Fix failed for patch {patch}: {e}")

        # 导出为二进制流
        buffer = io.BytesIO()
        self._doc.save(buffer)
        return buffer.getvalue()

    def _apply_patch(self, patch: Patch):
        """应用单个补丁详情到 docx 对象上。"""
        # 段落索引是 1-based
        if patch.para_index < 1 or patch.para_index > len(self._all_paras):
            return
            
        p_obj = self._all_paras[patch.para_index - 1]

        if patch.target_type == "run":
            if patch.run_index is not None and patch.run_index < len(p_obj.runs):
                r_obj = p_obj.runs[patch.run_index]
                self._fix_run(r_obj, patch.attribute, patch.value)
                
        elif patch.target_type == "paragraph":
            self._fix_para(p_obj, patch.attribute, patch.value)
            
        elif patch.target_type == "section":
            # 暂时只针对单节文档应用第一节配置（简化版 P1 逻辑）
            if len(self._doc.sections) > 0:
                s_obj = self._doc.sections[0]
                self._fix_section(s_obj, patch.attribute, patch.value)

    def _fix_run(self, run, attr, value):
        """
        物理修改 Run 的字体属性。
        """
        if attr == "bold":
            run.bold = bool(value)
        elif attr == "font_size":
            run.font.size = Pt(float(value))
        elif attr == "ascii_font":
            run.font.name = str(value)
        elif attr == "east_asia_font":
            # 中文字体需要特殊处理 XML EastAsia 属性 (w:eastAsia)
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.append(rFonts)
            rFonts.set(qn('w:eastAsia'), str(value))
            # 同时推荐将 ascii 也设为一致，防止回退渲染失败
            run.font.name = str(value)

    def _fix_para(self, para, attr, value):
        """
        物理修改段落排版属性。
        """
        pf = para.paragraph_format
        if attr == "alignment":
            # 映射对齐字符串到枚举
            align_map = {
                "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
                "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
                "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
                "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
            }
            para.alignment = align_map.get(str(value).lower())
        elif attr == "line_spacing":
            pf.line_spacing = float(value)
        elif attr == "first_line_indent_chars":
            # 这里简单处理，假设中文字体为 12pt（字号12pt * 2 = 24pt）
            # P2 改进：动态计算字体宽度的 indent
            pf.first_line_indent = Pt(float(value) * 12.0)
        elif attr == "space_before_pt":
            pf.space_before = Pt(float(value))
        elif attr == "space_after_pt":
            pf.space_after = Pt(float(value))

    def _fix_section(self, section, attr, value):
        """物理修改页边距。"""
        if attr == "top_margin_cm":
            section.top_margin = Cm(float(value))
        elif attr == "bottom_margin_cm":
            section.bottom_margin = Cm(float(value))
        elif attr == "left_margin_cm":
            section.left_margin = Cm(float(value))
        elif attr == "right_margin_cm":
            section.right_margin = Cm(float(value))
